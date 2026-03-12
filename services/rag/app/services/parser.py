"""Format-aware document text extraction.

Detects file type from extension and extracts plain text
using the appropriate library.
"""

import csv
import io
import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(data: bytes, filename: str) -> str:
    """Extract plain text from a document based on its file extension."""
    ext = Path(filename).suffix.lower()

    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".doc": _extract_docx,  # python-docx handles .doc if it's actually docx
        ".xlsx": _extract_xlsx,
        ".xls": _extract_xls,
        ".csv": _extract_csv,
        ".json": _extract_json,
    }

    extractor = extractors.get(ext)
    if extractor:
        try:
            return extractor(data)
        except Exception as e:
            logger.warning(f"Failed to parse {filename} as {ext}: {e}. Falling back to UTF-8.")
            return _extract_plaintext(data)

    return _extract_plaintext(data)


def _extract_pdf(data: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz)."""
    import fitz  # pymupdf

    pages = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text()
            if text.strip():
                pages.append(text.strip())

    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    """Extract text from DOCX using python-docx (paragraphs + tables)."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    """Extract text from XLSX using openpyxl."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        parts.append(f"## {sheet_name}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))

    wb.close()
    return "\n".join(parts)


def _extract_xls(data: bytes) -> str:
    """Extract text from legacy XLS using xlrd."""
    import xlrd

    wb = xlrd.open_workbook(file_contents=data)
    parts = []

    for sheet in wb.sheets():
        parts.append(f"## {sheet.name}")
        for row_idx in range(sheet.nrows):
            cells = [str(sheet.cell_value(row_idx, col)).strip()
                     for col in range(sheet.ncols)
                     if sheet.cell_value(row_idx, col)]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_csv(data: bytes) -> str:
    """Extract text from CSV."""
    text = data.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        cells = [c.strip() for c in row if c.strip()]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _extract_json(data: bytes) -> str:
    """Extract text from JSON — pretty-print or extract string values."""
    text = data.decode("utf-8", errors="ignore")
    try:
        obj = json.loads(text)
        return _json_to_text(obj)
    except json.JSONDecodeError:
        return text


def _json_to_text(obj: object, depth: int = 0) -> str:
    """Recursively extract readable text from JSON structures."""
    if isinstance(obj, str):
        return obj
    if isinstance(obj, (int, float, bool)):
        return str(obj)
    if isinstance(obj, list):
        parts = [_json_to_text(item, depth + 1) for item in obj]
        return "\n".join(p for p in parts if p)
    if isinstance(obj, dict):
        parts = []
        for key, value in obj.items():
            val_text = _json_to_text(value, depth + 1)
            if val_text:
                parts.append(f"{key}: {val_text}")
        return "\n".join(parts)
    return ""


def _extract_plaintext(data: bytes) -> str:
    """Fallback: decode as UTF-8."""
    return data.decode("utf-8", errors="ignore")
